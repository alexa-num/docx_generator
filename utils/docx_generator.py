"""Работа с DOCX файлами: замена плейсхолдеров и добавление приложений"""

import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm, Mm
import os
from PIL import Image


class DocxGenerator:
    """Класс для генерации DOCX документов"""
    
    def __init__(self, template_path=None):
        """
        Инициализация генератора
        
        Args:
            template_path: путь к файлу шаблона (если None, ищет в папке templates)
        """
        if template_path is None:
            # Ищем шаблон в папке templates
            base_dir = os.path.dirname(os.path.dirname(__file__))
            self.template_path = os.path.join(base_dir, 'templates', 'template.docx')
        else:
            self.template_path = template_path
    
    def check_template_exists(self):
        """Проверяет существует ли файл шаблона"""
        return os.path.exists(self.template_path)
    
    def get_template_path(self):
        """Возвращает путь к шаблону"""
        return self.template_path
    
    def generate_letter(self, output_path, letter_data, appendices, logo_path=None):
        """
        Генерирует письмо на основе шаблона и данных
        
        Args:
            output_path: путь для сохранения результата
            letter_data: словарь с данными письма
            appendices: список объектов Appendix
            logo_path: путь к файлу логотипа (опционально)
            
        Returns:
            bool: успешность операции
        """
        try:
            # Копируем шаблон
            shutil.copy2(self.template_path, output_path)
            doc = Document(output_path)
            
            # Сначала вставляем логотип (если есть)
            if logo_path and os.path.exists(logo_path):
                self._insert_logo(doc, logo_path)
            
            # Подготавливаем замены для остальных плейсхолдеров
            replacements = self._prepare_replacements(letter_data, appendices)
            
            # Заменяем все оставшиеся плейсхолдеры
            self._replace_all_placeholders(doc, replacements)
            
            # Добавляем приложения
            if appendices:
                self._add_appendices_to_doc(doc, appendices)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            raise Exception(f"Ошибка при генерации документа: {str(e)}")
    
    def _prepare_replacements(self, letter_data, appendices):
        """Подготавливает словарь замен для плейсхолдеров"""
        # Формируем список приложений
        appendices_list = self._generate_appendix_list(appendices)
        
        replacements = {
            "{COMPANY_LOGO}": "",  # Этот плейсхолдер уже удален при вставке логотипа
            "{COMPANY_INFO}": letter_data.get('company_info', ''),
            "{TO_JOB_TITLE}": letter_data.get('to_job_title', ''),
            "{TO_COMPANY}": letter_data.get('to_company', ''),
            "{TO_SHORT_NAME}": letter_data.get('to_short_name', ''),
            "{SUBJECT}": letter_data.get('subject', ''),
            "{TEXT}": letter_data.get('text', ''),
            "{FOOTER}": letter_data.get('footer', ''),
            "{FROM_TITLE}": letter_data.get('from_title', ''),
            "{FROM_NAME}": letter_data.get('from_name', ''),
            "{APPENDICES_LIST}": appendices_list
        }
        
        return replacements
    
    def _insert_logo(self, doc, logo_path):
        """
        Вставляет логотип в документ вместо плейсхолдера {COMPANY_LOGO}
        """
        try:
            max_width = Mm(40)  # 40 мм ширина
            max_height = Mm(15)  # 15 мм высота
            
            # Перебираем все параграфы в документе
            logo_found = False
            for paragraph in doc.paragraphs:
                # Проверяем наличие плейсхолдера в тексте параграфа
                if '{COMPANY_LOGO}' in paragraph.text:
                    # Очищаем параграф
                    paragraph.clear()
                    
                    # Создаем run для вставки изображения
                    run = paragraph.add_run()
                    
                    # Вставляем изображение
                    try:
                        run.add_picture(logo_path, width=max_width, height=max_height)
                        # Выравниваем по левому краю
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        logo_found = True
                        break
                    except Exception as e:
                        print(f"Ошибка при вставке изображения: {e}")
                        paragraph.text = "ООО 'ИНЕКСБАНК'"
                        logo_found = True
                        break
            
            # Если не нашли плейсхолдер, ищем в runs
            if not logo_found:
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if '{COMPANY_LOGO}' in run.text:
                            run.text = run.text.replace('{COMPANY_LOGO}', '')
                            run.add_picture(logo_path, width=max_width, height=max_height)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            logo_found = True
                            break
                    if logo_found:
                        break
            
            # Если все еще не нашли, вставляем перед COMPANY_INFO
            if not logo_found:
                for i, paragraph in enumerate(doc.paragraphs):
                    if '{COMPANY_INFO}' in paragraph.text:
                        new_paragraph = paragraph.insert_paragraph_before()
                        run = new_paragraph.add_run()
                        run.add_picture(logo_path, width=max_width, height=max_height)
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        logo_found = True
                        break
                
                    
            return True
            
        except Exception as e:
            print(f"Критическая ошибка при вставке логотипа: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _replace_all_placeholders(self, doc, replacements):
        """Заменяет все плейсхолдеры в документе"""
        
        # Замена в параграфах
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # Замена в колонтитулах
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
    
    def _replace_in_paragraph(self, paragraph, replacements):
        """Заменяет плейсхолдеры в одном параграфе"""
        # Проверяем каждый плейсхолдер
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                # Получаем весь текст параграфа
                full_text = paragraph.text
                new_text = full_text.replace(placeholder, value)
                
                # Сохраняем форматирование
                if paragraph.runs:
                    # Сохраняем свойства первого run
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    bold = first_run.bold
                    italic = first_run.italic
                    
                    # Очищаем параграф
                    paragraph.clear()
                    
                    # Добавляем новый текст с сохранением форматирования
                    run = paragraph.add_run(new_text)
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = font_size
                    if bold:
                        run.bold = bold
                    if italic:
                        run.italic = italic
                else:
                    # Если нет runs, просто заменяем текст
                    paragraph.text = new_text
    
    def _generate_appendix_list(self, appendices):
        """Формирует перечень приложений"""
        if not appendices:
            return ""
        
        lines = ["Приложения:"]
        for idx, app in enumerate(appendices, start=1):
            pages = getattr(app, 'pages', '1') + " л."
            lines.append(f"{idx}. {app.title} — {pages}")
        
        return "\n".join(lines)
    
    def _add_appendices_to_doc(self, doc, appendices):
        """Добавляет страницы с приложениями"""
        # Добавляем разрыв страницы перед приложениями
        if len(appendices) > 0:
            doc.add_page_break()
        
        for idx, app in enumerate(appendices, start=1):
            # Название приложения справа
            if len(appendices) == 1:
                app_title = doc.add_paragraph("Приложение")
            else:
                app_title = doc.add_paragraph(f"Приложение {idx}")
            
            app_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in app_title.runs:
                run.bold = True
            
            # Заголовок приложения по центру
            title_para = doc.add_paragraph(app.title)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.bold = True
                run.font.size = Inches(0.16)
            
            # Текст приложения
            doc.add_paragraph(app.text)
            
            # Информация о количестве листов
            pages = getattr(app, 'pages', '1')
            pages_para = doc.add_paragraph(f"Количество листов: {pages}")
            pages_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Разрыв страницы между приложениями
            if idx < len(appendices):
                doc.add_page_break()